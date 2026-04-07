"""
FreeCAD tutorial ingestion pipeline.

Reads sources from tutorial_sources.yaml and indexes them into a separate
ChromaDB collection (freecad_tutorials / chroma_tutorials/) used by the
tutorial-RAG variant of the agent.

Run:
    python scripts/ingest_tutorials.py              # ingest all sources
    python scripts/ingest_tutorials.py --clear      # wipe and re-ingest
    python scripts/ingest_tutorials.py --type wiki  # only wiki pages
    python scripts/ingest_tutorials.py --type youtube
    python scripts/ingest_tutorials.py --type local
    python scripts/ingest_tutorials.py --type external
    python scripts/ingest_tutorials.py --add-pdf path/to/file.pdf
    python scripts/ingest_tutorials.py --add-docx path/to/file.docx
    python scripts/ingest_tutorials.py --add-md path/to/file.md
    python scripts/ingest_tutorials.py --add-url https://example.com/page
    python scripts/ingest_tutorials.py --add-youtube VIDEO_ID_OR_URL
"""

import argparse
import re
import sys
from pathlib import Path

import requests
import yaml
from bs4 import BeautifulSoup
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from markdownify import markdownify

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

PROJECT_ROOT = Path(__file__).parent.parent
CHROMA_PATH = str(PROJECT_ROOT / "chroma_tutorials")
COLLECTION_NAME = "freecad_tutorials"
EMBED_MODEL = "all-MiniLM-L6-v2"
SOURCES_FILE = PROJECT_ROOT / "tutorial_sources.yaml"

CHUNK_SIZE = 800
CHUNK_OVERLAP = 150


# ---------------------------------------------------------------------------
# Loaders
# ---------------------------------------------------------------------------

def _scrape_html_page(url: str, source_type: str = "wiki") -> Document | None:
    """Scrape an HTML page and convert to markdown. Works for wiki and external URLs."""
    try:
        resp = requests.get(url, timeout=15)
        resp.raise_for_status()
    except Exception as e:
        print(f"  SKIP {url}: {e}")
        return None

    soup = BeautifulSoup(resp.text, "html.parser")

    # Extract title
    title_tag = (
        soup.find("h1", {"id": "firstHeading"})  # MediaWiki
        or soup.find("h1")
        or soup.find("title")
    )
    title = title_tag.get_text(strip=True) if title_tag else url.split("/")[-1]

    # Extract main content — try MediaWiki first, fall back to body
    content_div = (
        soup.find("div", {"id": "mw-content-text"})
        or soup.find("main")
        or soup.find("article")
        or soup.find("body")
    )
    if not content_div:
        print(f"  SKIP {url}: no content found")
        return None

    # Strip nav boxes, TOC, edit links
    for tag in content_div.find_all(
        ["div", "table", "span"],
        class_=["noprint", "navbox", "toc", "mw-editsection"],
    ):
        tag.decompose()

    text = markdownify(str(content_div), heading_style="ATX", strip=["a"]).strip()
    if len(text) < 100:
        print(f"  SKIP {url}: content too short")
        return None

    return Document(
        page_content=text,
        metadata={"source": url, "title": title, "type": source_type},
    )


def _load_pdf(path: Path) -> list[Document]:
    """Load a PDF file, one Document per page."""
    try:
        from pypdf import PdfReader
    except ImportError:
        print("  ERROR: pypdf not installed. Run: pip install pypdf")
        return []

    docs = []
    try:
        reader = PdfReader(str(path))
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = text.strip()
            if len(text) < 50:
                continue
            docs.append(Document(
                page_content=text,
                metadata={"source": str(path), "title": path.stem, "type": "pdf", "page": i + 1},
            ))
        print(f"  + {path.name} ({len(docs)} pages)")
    except Exception as e:
        print(f"  SKIP {path}: {e}")
    return docs


def _load_docx(path: Path) -> list[Document]:
    """Load a Word (.docx) file as a single Document."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        print("  ERROR: python-docx not installed. Run: pip install python-docx")
        return []

    try:
        doc = DocxDocument(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        if len(text) < 50:
            print(f"  SKIP {path}: content too short")
            return []
        print(f"  + {path.name}")
        return [Document(
            page_content=text,
            metadata={"source": str(path), "title": path.stem, "type": "docx"},
        )]
    except Exception as e:
        print(f"  SKIP {path}: {e}")
        return []


def _load_md(path: Path) -> list[Document]:
    """Load a Markdown file as a single Document."""
    try:
        text = path.read_text(encoding="utf-8", errors="replace").strip()
        if len(text) < 50:
            print(f"  SKIP {path}: content too short")
            return []
        print(f"  + {path.name}")
        return [Document(
            page_content=text,
            metadata={"source": str(path), "title": path.stem, "type": "markdown"},
        )]
    except Exception as e:
        print(f"  SKIP {path}: {e}")
        return []


def _extract_video_id(url_or_id: str) -> str:
    """Extract YouTube video ID from a URL or return the ID as-is."""
    url_or_id = url_or_id.strip()
    # Already an 11-char ID
    if re.match(r'^[A-Za-z0-9_-]{11}$', url_or_id):
        return url_or_id
    # youtu.be/ID
    m = re.search(r'youtu\.be/([A-Za-z0-9_-]{11})', url_or_id)
    if m:
        return m.group(1)
    # youtube.com/watch?v=ID
    m = re.search(r'[?&]v=([A-Za-z0-9_-]{11})', url_or_id)
    if m:
        return m.group(1)
    raise ValueError(f"Cannot extract video ID from: {url_or_id}")


def _load_youtube(url_or_id: str) -> list[Document]:
    """Fetch YouTube transcript and return as a single Document."""
    try:
        from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound
    except ImportError:
        print("  ERROR: youtube-transcript-api not installed. Run: pip install youtube-transcript-api")
        return []

    try:
        video_id = _extract_video_id(url_or_id)
    except ValueError as e:
        print(f"  SKIP {url_or_id}: {e}")
        return []

    try:
        transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
        text = " ".join(t["text"] for t in transcript_list).strip()
        if len(text) < 50:
            print(f"  SKIP {video_id}: transcript too short")
            return []
        url = f"https://www.youtube.com/watch?v={video_id}"
        print(f"  + YouTube:{video_id} ({len(text)} chars)")
        return [Document(
            page_content=text,
            metadata={"source": url, "title": f"YouTube:{video_id}", "type": "youtube"},
        )]
    except (TranscriptsDisabled, NoTranscriptFound) as e:
        print(f"  SKIP {video_id}: {e}")
        return []
    except Exception as e:
        print(f"  SKIP {video_id}: {e}")
        return []


def _load_local_file(path_str: str) -> list[Document]:
    """Dispatch to the correct loader based on file extension."""
    path = Path(path_str)
    if not path.is_absolute():
        path = PROJECT_ROOT / path
    if not path.exists():
        print(f"  SKIP {path}: file not found")
        return []
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    elif suffix == ".docx":
        return _load_docx(path)
    elif suffix in (".md", ".markdown"):
        return _load_md(path)
    else:
        print(f"  SKIP {path}: unsupported file type '{suffix}' (use .pdf, .docx, .md)")
        return []


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Ingest FreeCAD tutorials into ChromaDB")
    parser.add_argument("--clear", action="store_true", help="Wipe collection before ingesting")
    parser.add_argument(
        "--type",
        choices=["wiki", "external", "local", "youtube"],
        help="Only ingest this source type (default: all)",
    )
    # Ad-hoc single-file/URL additions
    parser.add_argument("--add-pdf",     metavar="PATH",     help="Add a single PDF file")
    parser.add_argument("--add-docx",    metavar="PATH",     help="Add a single DOCX file")
    parser.add_argument("--add-md",      metavar="PATH",     help="Add a single Markdown file")
    parser.add_argument("--add-url",     metavar="URL",      help="Add a single external URL")
    parser.add_argument("--add-youtube", metavar="ID_OR_URL", help="Add a single YouTube video")
    args = parser.parse_args()

    # Load sources config
    if not SOURCES_FILE.exists():
        print(f"ERROR: {SOURCES_FILE} not found. Create it from tutorial_sources.yaml template.")
        sys.exit(1)

    with open(SOURCES_FILE) as f:
        sources = yaml.safe_load(f) or {}

    wiki_urls     = sources.get("wiki_pages", []) or []
    external_urls = sources.get("external_urls", []) or []
    local_files   = sources.get("local_files", []) or []
    youtube_ids   = sources.get("youtube_videos", []) or []

    # Ad-hoc additions (not written back to yaml — edit the file manually for persistence)
    if args.add_pdf:
        local_files.append(args.add_pdf)
    if args.add_docx:
        local_files.append(args.add_docx)
    if args.add_md:
        local_files.append(args.add_md)
    if args.add_url:
        external_urls.append(args.add_url)
    if args.add_youtube:
        youtube_ids.append(args.add_youtube)

    print("Loading embedding model (downloads once on first run)...")
    embeddings = HuggingFaceEmbeddings(model_name=EMBED_MODEL)
    vectorstore = Chroma(
        collection_name=COLLECTION_NAME,
        embedding_function=embeddings,
        persist_directory=CHROMA_PATH,
    )

    if args.clear:
        print("Clearing existing tutorial collection...")
        try:
            vectorstore._collection.delete(where={"source": {"$ne": ""}})
        except Exception:
            pass

    raw_docs: list[Document] = []
    run_type = args.type  # None means all

    if run_type in (None, "wiki") and wiki_urls:
        print(f"\nScraping {len(wiki_urls)} FreeCAD wiki tutorial pages...")
        for url in wiki_urls:
            doc = _scrape_html_page(url, source_type="wiki")
            if doc:
                raw_docs.append(doc)
                print(f"  + {doc.metadata['title']}")

    if run_type in (None, "external") and external_urls:
        print(f"\nScraping {len(external_urls)} external URL(s)...")
        for url in external_urls:
            doc = _scrape_html_page(url, source_type="external")
            if doc:
                raw_docs.append(doc)
                print(f"  + {doc.metadata['title']}")

    if run_type in (None, "local") and local_files:
        print(f"\nLoading {len(local_files)} local file(s)...")
        for path_str in local_files:
            raw_docs.extend(_load_local_file(path_str))

    if run_type in (None, "youtube") and youtube_ids:
        print(f"\nFetching {len(youtube_ids)} YouTube transcript(s)...")
        for vid in youtube_ids:
            raw_docs.extend(_load_youtube(vid))

    if not raw_docs:
        print("\nNo documents collected. Add sources to tutorial_sources.yaml and retry.")
        sys.exit(1)

    print(f"\nChunking {len(raw_docs)} document(s)...")
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", " ", ""],
    )
    chunks = splitter.split_documents(raw_docs)
    print(f"  -> {len(chunks)} chunks")

    print("Embedding and storing in ChromaDB...")
    batch_size = 100
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i : i + batch_size]
        vectorstore.add_documents(batch)
        print(f"  {min(i + batch_size, len(chunks))}/{len(chunks)}")

    total = vectorstore._collection.count()
    print(f"\nDone. Tutorial ChromaDB now contains {total} chunks.")
    print(f"Location: {CHROMA_PATH}")


if __name__ == "__main__":
    main()
